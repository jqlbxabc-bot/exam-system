#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""试卷管理系统 - 试卷识别与自动分类模块"""

import os
import re
import json
import requests
import subprocess
from datetime import datetime

class ExamRecognizer:
    """试卷识别器，自动提取试卷信息"""
    VALID_SUBJECTS = {'语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治'}
    VALID_GRADES = {'高一', '高二', '高三'}
    VALID_EXAM_TYPES = {'月考', '期中', '期末', '模拟', '高考真题', '竞赛'}
    
    # 学科关键词
    SUBJECT_KEYWORDS = {
        '语文': ['语文', '作文', '阅读理解', '古诗文', '文言文', '现代文', '作文题', '默写', '诗词'],
        '数学': ['数学', '函数', '方程', '几何', '代数', '概率', '导数', '三角', '向量', '数列', '不等式'],
        '英语': ['英语', 'English', '听力', '阅读', '完形', '作文', '语法', '词汇', '翻译'],
        '物理': ['物理', '力学', '电学', '电磁', '光学', '热学', '动能', '势能', '牛顿', '电路'],
        '化学': ['化学', '有机', '无机', '反应', '元素', '化合物', '氧化', '还原', '电解', '溶液'],
        '生物': ['生物', '细胞', '基因', '遗传', '进化', '生态', '蛋白质', 'DNA', 'RNA', '光合作用'],
        '历史': ['历史', '朝代', '改革', '革命', '战争', '条约', '文明', '古代', '近代', '现代'],
        '地理': ['地理', '气候', '地形', '洋流', '板块', '人口', '城市', '农业', '工业', '资源'],
        '政治': ['政治', '经济', '哲学', '法律', '制度', '权利', '义务', '民主', '法治', '市场']
    }
    
    # 年级关键词
    GRADE_KEYWORDS = {
        '高一': ['高一', '高中一年级', '必修一', '必修二', '必修1', '必修2'],
        '高二': ['高二', '高中二年级', '选择性必修', '选修'],
        '高三': ['高三', '高中三年级', '高考', '模拟', '冲刺']
    }
    
    # 考试类型关键词
    EXAM_TYPE_KEYWORDS = {
        '月考': ['月考', '月测', '阶段测试'],
        '期中': ['期中', '中期'],
        '期末': ['期末', '期末考试'],
        '模拟': ['模拟', '模考', '仿真'],
        '高考真题': ['高考', '全国卷', '新高考'],
        '竞赛': ['竞赛', '奥林匹克', '奥赛']
    }
    
    # 题型关键词
    QUESTION_TYPE_KEYWORDS = {
        '选择题': ['选择题', '选择', '单选', '多选', '不定项选择'],
        '填空题': ['填空题', '填空'],
        '解答题': ['解答题', '计算题', '证明题', '应用题', '简答题', '论述题'],
        '实验题': ['实验题', '实验'],
        '作文题': ['作文', '写作']
    }

    @staticmethod
    def normalize_text(text):
        """清理OCR常见空格，便于规则识别。"""
        if not text:
            return ""
        text = text.replace('\u3000', ' ')
        text = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[\u4e00-\u9fff])', '', text)
        text = re.sub(r'(?<=[\u4e00-\u9fff])\s+(?=[A-Za-z0-9])', '', text)
        text = re.sub(r'(?<=[A-Za-z0-9])\s+(?=[\u4e00-\u9fff])', '', text)
        text = re.sub(r'(\d)\s+(?=分|分钟|年|月|日)', r'\1', text)
        text = re.sub(r'(?<=第)\s+([一二三四五六七八九十])', r'\1', text)
        return text

    @staticmethod
    def parse_json_response(content):
        """从AI回复中提取JSON，兼容代码块和前后解释文字。"""
        if not content:
            return None

        text = content.strip()
        if text.startswith('```'):
            text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.IGNORECASE)
            text = re.sub(r'\s*```$', '', text)

        candidates = [text]
        match = re.search(r'\{.*\}', text, re.DOTALL)
        if match:
            candidates.append(match.group(0))

        for candidate in candidates:
            try:
                return json.loads(candidate.strip(), strict=False)
            except Exception:
                pass

        fixed = candidates[-1]
        fixed = fixed.replace('，', ',').replace('：', ':')
        fixed = re.sub(r',\s*([}\]])', r'\1', fixed)
        try:
            return json.loads(fixed.strip(), strict=False)
        except Exception as e:
            print(f"AI JSON解析失败: {e}; 原始返回: {content[:500]}")
            return None
    
    @staticmethod
    def extract_text_from_image(image_path):
        """从图片中提取文字（使用OCR）"""
        try:
            # 尝试使用pytesseract
            import pytesseract
            from PIL import Image
            
            # 设置tesseract路径（Windows）
            import os
            tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_cmd):
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            
            image = Image.open(image_path)
            # 使用中文+英文识别
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            if text and text.strip():
                return text
        except ImportError:
            print("pytesseract未安装")
        except Exception as e:
            print(f"OCR识别出错: {e}")
        
        try:
            # 尝试使用百度OCR API
            result = ExamRecognizer._baidu_ocr(image_path)
            if result:
                return result
        except:
            pass
        
        return ""
    
    @staticmethod
    def _baidu_ocr(image_path):
        """使用百度OCR API识别文字"""
        import base64
        import requests
        
        # 这里需要配置百度OCR API Key
        from database import get_config
        api_key = get_config('baidu_ocr_api_key', '')
        secret_key = get_config('baidu_ocr_secret_key', '')
        
        if not api_key or not secret_key:
            return ""
        
        # 获取access_token
        token_url = f"https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id={api_key}&client_secret={secret_key}"
        token_resp = requests.post(token_url)
        access_token = token_resp.json().get('access_token')
        
        # 识别图片
        with open(image_path, 'rb') as f:
            image_data = base64.b64encode(f.read()).decode()
        
        ocr_url = f"https://aip.baidubce.com/rest/2.0/ocr/v1/general_basic?access_token={access_token}"
        resp = requests.post(ocr_url, data={'image': image_data})
        
        if resp.status_code == 200:
            result = resp.json()
            words = [item['words'] for item in result.get('words_result', [])]
            return '\n'.join(words)
        
        return ""
    
    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """从PDF中提取文字"""
        text = ""
        
        # 方法1: 使用PyMuPDF提取文字
        try:
            import fitz
            doc = fitz.open(pdf_path)
            for page in doc:
                text += page.get_text()
            doc.close()
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception as e:
            print(f"PyMuPDF提取失败: {e}")
        
        # 方法2: 使用PyPDF2
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception as e:
            print(f"PyPDF2提取失败: {e}")
        
        # 方法3: 对扫描件PDF使用OCR
        try:
            import fitz
            import pytesseract
            from PIL import Image
            import tempfile
            
            # 设置tesseract路径
            tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
            if os.path.exists(tesseract_cmd):
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            
            doc = fitz.open(pdf_path)
            ocr_text = ""
            
            # 对每一页进行OCR（最多5页）
            max_pages = min(5, len(doc))
            for page_num in range(max_pages):
                page = doc[page_num]
                # 转换为图片
                zoom = 2.0
                mat = fitz.Matrix(zoom, zoom)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # 保存为临时文件
                temp_path = os.path.join(tempfile.gettempdir(), f'ocr_page_{page_num}.png')
                with open(temp_path, 'wb') as f:
                    f.write(img_data)
                
                # OCR识别
                image = Image.open(temp_path)
                page_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                ocr_text += page_text + "\n"
                
                # 清理临时文件
                os.remove(temp_path)
            
            doc.close()
            
            if ocr_text.strip():
                print(f"OCR识别成功，文字长度: {len(ocr_text)}")
                return ocr_text
                
        except Exception as e:
            print(f"OCR识别失败: {e}")
        
        return text
    
    @staticmethod
    def extract_text_from_docx(docx_path):
        """从Word文档中提取文字，包含表格内容。"""
        try:
            import docx
            doc = docx.Document(docx_path)
            parts = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except ImportError:
            print("python-docx未安装")
        except Exception as e:
            print(f"Word文档提取失败: {e}")

        return ""

    @staticmethod
    def extract_text_from_doc(doc_path):
        """尽量从旧版 .doc 文档提取文字。"""
        # 有些 .doc 实际是 docx 格式，先试 python-docx。
        text = ExamRecognizer.extract_text_from_docx(doc_path)
        if text.strip():
            return text

        # Windows 本地运行时，如果安装了 Word/pywin32，可以走 COM 自动化。
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(doc_path))
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            return text or ""
        except Exception as e:
            print(f"旧版Word(.doc)提取失败: {e}")

        for cmd in (['antiword', doc_path], ['catdoc', '-w', doc_path]):
            try:
                completed = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    encoding='utf-8',
                    errors='ignore',
                    timeout=30
                )
                if completed.returncode == 0 and completed.stdout.strip():
                    return completed.stdout
            except Exception as e:
                print(f".doc命令提取失败({cmd[0]}): {e}")
        
        return ""
    
    @staticmethod
    def extract_text(file_path):
        """根据文件类型提取文字"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return ExamRecognizer.extract_text_from_image(file_path)
        elif ext == '.pdf':
            return ExamRecognizer.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return ExamRecognizer.extract_text_from_docx(file_path)
        elif ext == '.doc':
            return ExamRecognizer.extract_text_from_doc(file_path)
        
        return ""
    
    @staticmethod
    def recognize_subject(text):
        """识别学科"""
        text = ExamRecognizer.normalize_text(text or "")
        text_lower = text.lower()
        header = text[:1200]
        header_lower = header.lower()
        scores = {}
        
        for subject, keywords in ExamRecognizer.SUBJECT_KEYWORDS.items():
            score = 0
            for keyword in keywords:
                keyword_lower = keyword.lower()
                if keyword_lower in header_lower:
                    score += 3
                elif keyword_lower in text_lower:
                    score += 1
            # 标题/页眉里的“科目+试卷/试题/考试”比正文知识点更可靠。
            if re.search(rf'{subject}.{{0,12}}(试卷|试题|考试|测试|检测|真题)', header):
                score += 8
            if re.search(rf'(试卷|试题|考试|测试|检测|真题).{{0,12}}{subject}', header):
                score += 8
            if score > 0:
                scores[subject] = score
        
        if scores:
            return max(scores, key=scores.get)
        return None
    
    @staticmethod
    def recognize_grade(text):
        """识别年级"""
        text = ExamRecognizer.normalize_text(text or "")
        direct_patterns = {
            '高一': [r'高一', r'高中一年级', r'高1'],
            '高二': [r'高二', r'高中二年级', r'高2'],
            '高三': [r'高三', r'高中三年级', r'高3'],
        }
        for grade, patterns in direct_patterns.items():
            if any(re.search(pattern, text) for pattern in patterns):
                return grade
        for grade, keywords in ExamRecognizer.GRADE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text:
                    return grade
        return None
    
    @staticmethod
    def recognize_exam_type(text):
        """识别考试类型"""
        text = ExamRecognizer.normalize_text(text or "")
        header = text[:1200]
        if re.search(r'(模拟|模考|仿真|一模|二模|三模)', header):
            return '模拟'
        if re.search(r'(期末|期末考试)', header):
            return '期末'
        if re.search(r'(期中|中期)', header):
            return '期中'
        if re.search(r'(月考|月测|阶段测试)', header):
            return '月考'
        if re.search(r'(竞赛|奥林匹克|奥赛)', header):
            return '竞赛'
        if re.search(r'(高考|全国卷|新高考).{0,12}(真题|试题|试卷)', header):
            return '高考真题'
        return None
    
    @staticmethod
    def recognize_total_score(text):
        """识别总分"""
        text = ExamRecognizer.normalize_text(text or "")
        patterns = [
            r'总分[：:]\s*(\d+)\s*分',
            r'满分[：:]\s*(\d+)\s*分',
            r'满分\s*(\d+)\s*分',
            r'全卷满分\s*(\d+)\s*分',
            r'本试卷.{0,30}满分\s*(\d+)\s*分',
            r'试卷.{0,30}满分\s*(\d+)\s*分',
            r'共\s*(\d+)\s*分',
            r'(\d+)\s*分\s*满分',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                value = float(match.group(1))
                if value >= 60 or '满分' in pattern or '总分' in pattern:
                    return value

        header = text[:1500]
        scores = [int(x) for x in re.findall(r'(?<!小题)(?<!每题)(?<!每小题)(\d{2,3})\s*分', header)]
        plausible = [score for score in scores if 60 <= score <= 300]
        if plausible:
            return float(max(plausible))
        
        return None
    
    @staticmethod
    def recognize_exam_date(text):
        """识别考试日期"""
        text = ExamRecognizer.normalize_text(text or "")
        patterns = [
            r'(\d{4})[年/-](\d{1,2})[月/-](\d{1,2})[日号]?',
            r'(\d{4})[年/-](\d{1,2})月',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                try:
                    year = int(match.group(1))
                    month = int(match.group(2))
                    day = int(match.group(3)) if len(match.groups()) > 2 else 1
                    if 2000 <= year <= 2100 and 1 <= month <= 12:
                        return f"{year}-{month:02d}-{day:02d}"
                except:
                    pass
        
        return None
    
    @staticmethod
    def extract_title(text, source_name=None):
        """提取试卷标题"""
        text = ExamRecognizer.normalize_text(text or "")
        noise_patterns = [
            '绝密', '启用前', '注意事项', '本试卷', '答题卡', '姓名', '准考证号',
            '考试时间', '满分', '总分', '班级', '学校'
        ]
        # 尝试从第一行提取标题
        lines = text.strip().split('\n')
        candidates = []
        for line in lines[:12]:
            line = re.sub(r'\s+', ' ', line.strip())
            if not line or any(noise in line for noise in noise_patterns):
                continue
            if len(line) > 5 and len(line) < 100:
                # 检查是否包含年份、科目等关键词
                if any(keyword in line for keyword in ['年', '学期', '考试', '测试', '试卷', '检测']):
                    candidates.append(line)
        if candidates:
            return candidates[0]
        
        if source_name:
            stem = os.path.splitext(os.path.basename(source_name))[0]
            stem = re.sub(r'[_\-]+', ' ', stem).strip()
            if stem:
                return stem[:80]
        
        # 如果没找到，返回前30个字符
        first_line = lines[0].strip() if lines else "未命名试卷"
        return first_line[:50] if len(first_line) > 50 else first_line
    
    @staticmethod
    def extract_questions(text):
        """提取题目 - 仅匹配行首的题号，避免误匹配文中数字"""
        questions = []
        if not text:
            return questions
        
        lines = text.strip().split('\n')
        
        # 逐行匹配题号模式
        question_start_pattern = re.compile(
            r'^[（(]?\s*(\d{1,3})\s*[)）]{1}\s*'   # (1) 或 1)
            r'|^\s*(\d{1,3})\s*[.、．]\s*'         # 1. 或 1、
            r'|^第\s*(\d{1,3})\s*题\s*'             # 第一题、第二题
        )
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            match = question_start_pattern.search(line)
            if match:
                groups = match.groups()
                # 三个捕获组中只有一个会匹配成功
                num_str = next((g for g in groups if g is not None), None)
                if num_str:
                    num = int(num_str)
                    # 截取题号后的内容
                    content_start = match.end()
                    content = line[content_start:].strip()
                    
                    # 收集后续行（直到遇到下一个题号或空行后接题号）
                    j = i + 1
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if not next_line:
                            # 空行，检查后面是否是新题号
                            if j + 1 < len(lines) and question_start_pattern.match(lines[j + 1].strip()):
                                break
                            j += 1
                            continue
                        if question_start_pattern.match(next_line):
                            break
                        content += '\n' + lines[j].strip()
                        j += 1
                    
                    i = j  # 跳过已处理的行
                    
                    if len(content.strip()) > 3:
                        questions.append({
                            'number': num,
                            'content': content.strip()[:500],
                            'type': ExamRecognizer._guess_question_type(content)
                        })
                    continue
            i += 1
        
        return questions
    
    @staticmethod
    def _guess_question_type(content):
        """猜测题型"""
        content_lower = content.lower()
        
        if any(keyword in content_lower for keyword in ['选择', 'a.', 'b.', 'c.', 'd.', 'A.', 'B.', 'C.', 'D.']):
            return '选择题'
        elif any(keyword in content_lower for keyword in ['填空', '___', '（  ）', '(  )']):
            return '填空题'
        elif any(keyword in content_lower for keyword in ['计算', '证明', '解答', '求', '解']):
            return '解答题'
        elif any(keyword in content_lower for keyword in ['实验', '操作', '观察']):
            return '实验题'
        elif any(keyword in content_lower for keyword in ['作文', '写作', '不少于']):
            return '作文题'
        
        return '解答题'
    
    @staticmethod
    def analyze_with_ai(text, file_paths=None, source_name=None):
        """使用AI分析试卷内容"""
        from ai_analyzer import get_analyzer
        
        prompt = f"""请分析以下试卷内容，提取关键信息并以JSON格式返回：

文件名：
{source_name or '未知'}

试卷内容：
{text[:3000]}

请返回以下格式的JSON：
{{
    "title": "试卷标题",
    "subject": "学科（语文/数学/英语/物理/化学/生物/历史/地理/政治）",
    "grade": "年级（高一/高二/高三）",
    "exam_type": "考试类型（月考/期中/期末/模拟/高考真题/竞赛）",
    "total_score": 总分（数字）,
    "exam_date": "考试日期（YYYY-MM-DD格式，如果无法识别则为null）",
    "description": "试卷简要描述",
    "knowledge_points": ["知识点1", "知识点2"],
    "difficulty": "难度（简单/中等/困难）",
    "question_types": {{
        "选择题": 数量,
        "填空题": 数量,
        "解答题": 数量
    }}
}}

只返回JSON，不要其他内容。"""
        
        try:
            analyzer = get_analyzer()
            result = analyzer.chat(prompt)
            return ExamRecognizer.parse_json_response(result)
        except Exception as e:
            print(f"AI分析调用失败: {e}")
            return None
    
    @staticmethod
    def _clean_ai_result(ai_result):
        """过滤AI返回值，避免无效字段覆盖基础识别结果。"""
        if not isinstance(ai_result, dict):
            return {}

        cleaned = {}
        subject = ai_result.get('subject')
        if subject in ExamRecognizer.VALID_SUBJECTS:
            cleaned['subject'] = subject

        grade = ai_result.get('grade')
        if grade in ExamRecognizer.VALID_GRADES:
            cleaned['grade'] = grade

        exam_type = ai_result.get('exam_type')
        if exam_type in ExamRecognizer.VALID_EXAM_TYPES:
            cleaned['exam_type'] = exam_type

        title = ai_result.get('title')
        if isinstance(title, str) and 3 <= len(title.strip()) <= 120:
            cleaned['title'] = title.strip()

        description = ai_result.get('description')
        if isinstance(description, str) and description.strip():
            cleaned['description'] = description.strip()[:300]

        total_score = ai_result.get('total_score')
        try:
            total_score = float(total_score)
            if 1 <= total_score <= 1000:
                cleaned['total_score'] = total_score
        except (TypeError, ValueError):
            pass

        exam_date = ai_result.get('exam_date')
        if isinstance(exam_date, str) and re.match(r'^\d{4}-\d{2}-\d{2}$', exam_date):
            cleaned['exam_date'] = exam_date

        return cleaned

    @staticmethod
    def recognize_exam(file_path, use_ai=True, original_filename=None):
        """综合识别试卷"""
        result = {
            'title': None,
            'subject': None,
            'grade': None,
            'exam_type': None,
            'total_score': None,
            'exam_date': None,
            'description': None,
            'text': None,
            'questions': [],
            'ai_analysis': None
        }
        
        # 判断文件类型
        ext = os.path.splitext(file_path)[1].lower()
        is_image = ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
        is_pdf = ext == '.pdf'
        
        print(f"开始识别试卷: {file_path}, 类型: {ext}")
        
        # 尝试OCR提取文字
        text = ""
        try:
            text = ExamRecognizer.extract_text(file_path)
            print(f"OCR提取文字长度: {len(text) if text else 0}")
        except Exception as e:
            print(f"OCR提取失败: {e}")
        
        filename_context = os.path.splitext(os.path.basename(original_filename or file_path))[0]
        normalized_text = ExamRecognizer.normalize_text(text)
        recognition_text = f"{filename_context}\n{normalized_text}" if filename_context else normalized_text
        result['text'] = text
        
        # 如果有文字，进行基础识别
        if recognition_text and len(recognition_text) > 2:
            result['title'] = ExamRecognizer.extract_title(normalized_text, original_filename)
            result['subject'] = ExamRecognizer.recognize_subject(recognition_text)
            result['grade'] = ExamRecognizer.recognize_grade(recognition_text)
            result['exam_type'] = ExamRecognizer.recognize_exam_type(text) or ExamRecognizer.recognize_exam_type(filename_context)
            result['total_score'] = ExamRecognizer.recognize_total_score(recognition_text)
            result['exam_date'] = ExamRecognizer.recognize_exam_date(recognition_text)
            result['questions'] = ExamRecognizer.extract_questions(text)
            print(f"基础识别结果: 标题={result['title']}, 科目={result['subject']}")
        
        # 使用AI进行识别
        if use_ai:
            try:
                if is_image:
                    # 图片：优先AI视觉识别（直接看懂图片内容），失败再降级OCR
                    print("尝试AI图片视觉分析...")
                    ai_result = ExamRecognizer.analyze_image_with_ai(file_path)
                    if not ai_result and text and len(text) > 10:
                        print("视觉识别失败/不支持，降级为OCR文本分析...")
                        ai_result = ExamRecognizer.analyze_with_ai(text, [file_path], original_filename)
                    elif not ai_result:
                        print("图片AI识别完全失败（无视觉模型且OCR无文字）")
                elif is_pdf:
                    # PDF文件，用AI分析PDF（会内部提取文字+调用AI）
                    print("使用AI分析PDF...")
                    ai_result = ExamRecognizer.analyze_pdf_with_ai(file_path)
                elif text and len(text) > 10:
                    # 有文字内容的文档：用AI分析文字
                    print("使用AI分析OCR文字...")
                    ai_result = ExamRecognizer.analyze_with_ai(text, [file_path], original_filename)
                else:
                    # docx/doc 等文档格式：OCR 可能已失败，直接尝试用文档转文字+AI
                    print(f"OCR未提取文字，使用AI分析文档内容...")
                    doc_text = ExamRecognizer.extract_text(file_path)
                    if doc_text and len(doc_text) > 10:
                        ai_result = ExamRecognizer.analyze_with_ai(doc_text, [file_path], original_filename)
                    else:
                        ai_result = None
                
                if ai_result:
                    result['ai_analysis'] = ai_result
                    ai_result = ExamRecognizer._clean_ai_result(ai_result)
                    # 用AI结果补充或覆盖基础识别结果
                    if ai_result.get('title'):
                        result['title'] = ai_result['title']
                    if ai_result.get('subject'):
                        result['subject'] = ai_result['subject']
                    if ai_result.get('grade'):
                        result['grade'] = ai_result['grade']
                    if ai_result.get('exam_type'):
                        result['exam_type'] = ai_result['exam_type']
                    if ai_result.get('total_score'):
                        result['total_score'] = float(ai_result['total_score'])
                    if ai_result.get('exam_date'):
                        result['exam_date'] = ai_result['exam_date']
                    if ai_result.get('description'):
                        result['description'] = ai_result['description']
                    print(f"AI识别结果: 标题={result['title']}, 科目={result['subject']}")
            except Exception as e:
                print(f"AI识别失败: {e}")
        
        return result
    
    @staticmethod
    def analyze_image_with_ai(image_path):
        """使用AI视觉模型分析图片。
        优先使用视觉模型识别；不支持视觉时自动降级为 OCR + 纯文本分析。"""
        from ai_analyzer import get_analyzer
        import base64

        analyzer = get_analyzer()

        # 已知支持视觉的模型列表（避免文本模型收到图片请求后 400 错误）
        vision_model_keywords = ['vision', 'vl', 'vl2', 'gpt-4o', 'gpt-4v', 'claude-3',
                                  'gemini', 'qwen-vl', 'glm-4v', 'yi-vision']
        model_lower = (analyzer.model or '').lower()
        is_vision_model = any(kw in model_lower for kw in vision_model_keywords)
        # 也检查通用的 vision-compatible provider（如果模型名不明确但 provider 已知支持 vision）
        vision_providers = ['openai', 'moonshot', 'qwen', 'zhipu']

        use_vision = is_vision_model and analyzer.api_key

        if not use_vision and not is_vision_model and analyzer.provider in vision_providers:
            # provider 支持 vision 但模型名无法确认，则尝试发送（可能会失败回退）
            use_vision = True

        prompt_text = """你是一个专业的试卷分析助手。请仔细查看这张试卷图片，识别图片中所有的文字内容，包括：

- 试卷标题（通常在最上方，大字）
- 考试说明和注意事项
- 每一道题目的题号、题干内容
- 选择题的选项（A/B/C/D）
- 如果有答案，也一并提取

请根据识别到的内容，以JSON格式返回：

{
    "title": "试卷标题（从图片顶部大字提取）",
    "subject": "学科（根据实际题目内容判断：语文/数学/英语/物理/化学/生物/历史/地理/政治）",
    "grade": "年级（根据题目难度判断：高一/高二/高三）",
    "exam_type": "考试类型（月考/期中/期末/模拟/高考真题/竞赛）",
    "total_score": 总分数字（如150、100）,
    "exam_date": "YYYY-MM-DD（如无法识别则为null）",
    "description": "简要描述试卷结构和内容（如：包含选择题12道、填空题4道、解答题6道）",
    "text": "从图片中提取的前500字实际文字内容"
}

重要：text字段必须包含从图片中实际读到的文字，不能是模板文字或占位符。只返回JSON，不要其他内容。"""

        if not use_vision or not analyzer.api_key:
            return None

        try:
            with open(image_path, 'rb') as f:
                img_data = base64.b64encode(f.read()).decode()

            # 根据文件扩展名确定 MIME 类型
            ext = os.path.splitext(image_path)[1].lower()
            mime_map = {'.jpg': 'jpeg', '.jpeg': 'jpeg', '.png': 'png',
                       '.gif': 'gif', '.bmp': 'bmp', '.webp': 'webp'}
            mime_type = mime_map.get(ext, 'jpeg')

            messages = [{
                'role': 'user',
                'content': [
                    {'type': 'text', 'text': prompt_text},
                    {'type': 'image_url',
                     'image_url': {'url': f'data:image/{mime_type};base64,{img_data}'}}
                ]
            }]
            headers = {
                'Authorization': f'Bearer {analyzer.api_key}',
                'Content-Type': 'application/json'
            }
            data = {
                'model': analyzer.model,
                'messages': messages,
                'max_tokens': 4000,
                'temperature': 0.3
            }
            response = requests.post(
                f'{analyzer.base_url}/chat/completions',
                headers=headers,
                json=data,
                timeout=120
            )
            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                parsed = ExamRecognizer.parse_json_response(content)
                if parsed and parsed.get('subject') and parsed.get('title'):
                    return parsed
                print(f"视觉识别返回内容不完整: {parsed}")
                return parsed

            # 400 通常表示模型不支持 vision
            if response.status_code == 400:
                print(f"视觉识别失败(模型可能不支持图片): {response.text[:200]}")
            else:
                print(f"视觉识别 API 返回 {response.status_code}")

        except Exception as e:
            print(f"AI视觉分析失败: {e}")

        # 降级：OCR 提取文字后用文本 AI 分析
        try:
            ocr_text = ExamRecognizer.extract_text(image_path)
            if ocr_text and len(ocr_text.strip()) > 20:
                print(f"视觉识别失败，使用OCR文字({len(ocr_text)}字符)进行文本分析...")
                source_name = os.path.basename(image_path)
                return ExamRecognizer.analyze_with_ai(ocr_text, [image_path], source_name)
        except Exception as e2:
            print(f"OCR 降级方案也失败: {e2}")

        return None
    
    @staticmethod
    def analyze_pdf_with_ai(pdf_path):
        """使用AI分析PDF文件"""
        import fitz  # PyMuPDF
        
        try:
            from ai_analyzer import get_analyzer
            analyzer = get_analyzer()
            
            # 打开PDF提取文字
            doc = fitz.open(pdf_path)
            full_text = ""
            for page in doc:
                full_text += page.get_text()
            doc.close()
            
            if not full_text or len(full_text.strip()) < 10:
                print("PDF文字提取失败，尝试 OCR 提取...")
                full_text = ExamRecognizer.extract_text_from_pdf(pdf_path)
            
            if not full_text or len(full_text.strip()) < 10:
                print("PDF文字和OCR提取均失败，内容太少")
                return None
            
            print(f"PDF提取文字长度: {len(full_text)}")
            
            # 使用AI分析文字内容
            prompt = f"""请分析以下试卷内容，提取信息并以JSON格式返回：

试卷内容：
{full_text[:3000]}

请提取以下信息：
1. title: 试卷标题
2. subject: 学科（语文/数学/英语/物理/化学/生物/历史/地理/政治）
3. grade: 年级（高一/高二/高三）
4. exam_type: 考试类型（月考/期中/期末/模拟/高考真题/竞赛）
5. total_score: 总分（数字）
6. exam_date: 考试日期（YYYY-MM-DD格式，如果无法识别则为null）
7. description: 试卷简要描述

只返回JSON，不要其他内容。"""

            result = analyzer.chat(prompt)
            return ExamRecognizer.parse_json_response(result)
            
        except Exception as e:
            print(f"AI PDF分析失败: {e}")
            return None


def auto_recognize_and_save(file_path, user_id=None):
    """自动识别并保存试卷"""
    from database import add_exam, add_question, add_ai_analysis
    
    # 识别试卷
    result = ExamRecognizer.recognize_exam(file_path, use_ai=True)
    
    # 保存到数据库
    exam_id = add_exam(
        title=result.get('title') or '未命名试卷',
        subject=result.get('subject') or '其他',
        grade=result.get('grade'),
        exam_type=result.get('exam_type') or '其他',
        exam_date=result.get('exam_date'),
        total_score=result.get('total_score') or 150,
        description=result.get('description'),
        file_path=file_path,
        file_type=os.path.splitext(file_path)[1][1:],
        upload_user_id=user_id
    )
    
    # 保存题目
    for q in result.get('questions', []):
        add_question(
            exam_id=exam_id,
            question_number=q.get('number'),
            question_type=q.get('type', '解答题'),
            content=q.get('content')
        )
    
    # 保存AI分析结果
    if result.get('ai_analysis'):
        add_ai_analysis(
            exam_id=exam_id,
            analysis_type='自动识别',
            content=json.dumps(result['ai_analysis'], ensure_ascii=False, indent=2),
            model_name='auto'
        )
    
    return {
        'exam_id': exam_id,
        'recognition_result': result
    }


if __name__ == '__main__':
    # 测试
    recognizer = ExamRecognizer()
    
    test_text = """
    2026年高一下学期数学期中考试试卷
    总分：150分  时间：120分钟
    
    一、选择题（每小题5分，共60分）
    1. 已知集合A={1,2,3}，B={2,3,4}，则A∩B=
    A. {1,2}  B. {2,3}  C. {3,4}  D. {1,4}
    
    二、填空题（每小题5分，共20分）
    13. 函数f(x)=x²+2x+1的最小值为______
    
    三、解答题（共70分）
    17. （10分）已知等差数列{an}中，a1=2，a3=6，求通项公式。
    """
    
    print("学科识别:", recognizer.recognize_subject(test_text))
    print("年级识别:", recognizer.recognize_grade(test_text))
    print("考试类型:", recognizer.recognize_exam_type(test_text))
    print("总分识别:", recognizer.recognize_total_score(test_text))
    print("标题提取:", recognizer.extract_title(test_text))
