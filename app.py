#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""试卷管理系统 - 主程序"""

from flask import (Flask, render_template, request, redirect, url_for, 
                   session, jsonify, flash, send_from_directory)
from werkzeug.utils import secure_filename
import os
import json
from datetime import datetime, timedelta

import database as db
from cache import cache, cached, invalidate_cache, CACHE_KEYS
from ai_analyzer import get_analyzer
from storage import get_storage_manager, allowed_file
from gaokao_classifier import (
    get_classification_prompt, parse_classification_result,
    classify_question_type, estimate_difficulty, SUBJECT_CATEGORIES
)

app = Flask(__name__)
app.secret_key = 'exam_system_secret_key_2026'
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

# 初始化
db.init_db()

# ==================== 辅助函数 ====================
def login_required(f):
    """登录验证装饰器"""
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

def get_current_user():
    """获取当前用户"""
    return session.get('user')

# ==================== 路由 ====================
@app.route('/')
def index():
    if 'user' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = db.verify_user(username, password)
        if user:
            session['user'] = user
            return redirect(url_for('dashboard'))
        flash('用户名或密码错误', 'error')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    user = get_current_user()
    recent_exams = db.get_all_exams(limit=10)
    
    # 统计数据 - 使用缓存
    cache_key = f"dashboard:stats:{user['id']}"
    stats = cache.get(cache_key)
    
    if stats is None:
        all_exams = db.get_all_exams(limit=10000)
        mistakes = db.get_mistakes(user['id']) if user['role'] == 'student' else []
        
        # 计算正确率
        total_questions = 0
        correct_questions = 0
        for exam in all_exams:
            questions = db.get_questions(exam['id'])
            total_questions += len(questions)
            correct_questions += sum(1 for q in questions if q.get('user_score') and q['user_score'] > 0)
        
        accuracy_rate = round(correct_questions / total_questions * 100, 1) if total_questions > 0 else 0
        
        stats = {
            'total_exams': len(all_exams),
            'total_questions': total_questions,
            'total_mistakes': len(mistakes),
            'accuracy_rate': accuracy_rate,
        }
        
        # 缓存5分钟
        cache.set(cache_key, stats, timeout=300)
    
    # 成绩趋势数据（最近10次考试）
    study_stats = db.get_study_stats(user['id']) if user['role'] == 'student' else []
    score_dates = [s['exam_date'][-5:] for s in study_stats[-10:]]  # 取月-日
    score_values = [round(s['score'] / s['total_score'] * 100, 1) for s in study_stats[-10:]]
    
    # 科目分布数据 - 使用缓存
    subject_cache_key = "dashboard:subject_dist"
    subject_data = cache.get(subject_cache_key)
    
    if subject_data is None:
        all_exams = db.get_all_exams(limit=10000)
        subject_counts = {}
        for exam in all_exams:
            subject = exam.get('subject', '未知')
            subject_counts[subject] = subject_counts.get(subject, 0) + 1
        subject_data = {
            'names': list(subject_counts.keys()),
            'values': list(subject_counts.values())
        }
        # 缓存10分钟
        cache.set(subject_cache_key, subject_data, timeout=600)
    
    # 需要复习的错题（今天或已过期）
    from datetime import date
    today = date.today().isoformat()
    mistakes = db.get_mistakes(user['id']) if user['role'] == 'student' else []
    review_mistakes = [m for m in mistakes if m.get('next_review_date') and m['next_review_date'] <= today]
    
    # 最近活动
    recent_activities = []
    for exam in recent_exams[:5]:
        recent_activities.append({
            'description': f'上传了试卷《{exam["title"]}》',
            'time': exam['created_at'][:16]
        })
    
    return render_template('dashboard.html', 
                          user=user, 
                          exams=recent_exams, 
                          stats=stats,
                          score_dates=score_dates,
                          score_values=score_values,
                          subject_names=subject_data['names'],
                          subject_counts=subject_data['values'],
                          review_mistakes=review_mistakes,
                          recent_activities=recent_activities)

# ==================== 试卷管理 ====================
@app.route('/exams')
@login_required
def exam_list():
    keyword = request.args.get('keyword', '')
    subject = request.args.get('subject', '')
    grade = request.args.get('grade', '')
    exam_type = request.args.get('exam_type', '')
    page = request.args.get('page', 1, type=int)
    
    # 使用分页查询
    result = db.search_exams(keyword=keyword, subject=subject, grade=grade, 
                            exam_type=exam_type, page=page, per_page=20)
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    grades = ['高一', '高二', '高三']
    exam_types = ['月考', '期中', '期末', '模拟', '高考真题', '竞赛']
    
    return render_template('exam_list.html', 
                          exams=result['items'],
                          pagination=result,
                          subjects=subjects, grades=grades, exam_types=exam_types,
                          keyword=keyword, subject=subject, grade=grade, exam_type=exam_type)

@app.route('/exams/batch-import')
@login_required
def batch_import():
    """批量导入页面"""
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    grades = ['高一', '高二', '高三']
    return render_template('batch_import.html', subjects=subjects, grades=grades)

@app.route('/exams/upload', methods=['GET', 'POST'])
@login_required
def upload_exam():
    if request.method == 'POST':
        try:
            title = request.form.get('title', '').strip()
            subject = request.form.get('subject', '').strip()
            grade = request.form.get('grade', '').strip()
            exam_type = request.form.get('exam_type', '月考').strip()
            exam_date = request.form.get('exam_date', '').strip()
            total_score = request.form.get('total_score', '150')
            description = request.form.get('description', '').strip()
            auto_recognize = request.form.get('auto_recognize') == 'on'
            
            try:
                total_score = float(total_score)
            except:
                total_score = 150.0
            
            user = get_current_user()
            
            # 处理文件上传
            file_path = None
            file_type = None
            cloud_url = None
            
            if 'file' in request.files:
                file = request.files['file']
                if file and file.filename and allowed_file(file.filename):
                    storage = get_storage_manager()
                    upload_result = storage.upload_file(file, folder='exams')
                    if upload_result['success']:
                        file_path = upload_result.get('file_path')
                        file_type = file.filename.rsplit('.', 1)[1].lower()
                        cloud_url = upload_result.get('cloud_url')
            
            # 自动识别试卷
            recognition_result = None
            if auto_recognize and file_path:
                try:
                    from exam_recognizer import ExamRecognizer
                    recognizer = ExamRecognizer()
                    # 优先使用AI识别（包括AI视觉识别）
                    recognition_result = recognizer.recognize_exam(file_path, use_ai=True)
                    
                    # 用识别结果填充空字段
                    if not title and recognition_result.get('title'):
                        title = recognition_result['title']
                    if not subject and recognition_result.get('subject'):
                        subject = recognition_result['subject']
                    if not grade and recognition_result.get('grade'):
                        grade = recognition_result['grade']
                    if not exam_type and recognition_result.get('exam_type'):
                        exam_type = recognition_result['exam_type']
                    if not exam_date and recognition_result.get('exam_date'):
                        exam_date = recognition_result['exam_date']
                    if recognition_result.get('total_score'):
                        total_score = recognition_result['total_score']
                    if not description and recognition_result.get('description'):
                        description = recognition_result['description']
                except Exception as e:
                    print(f"自动识别失败: {e}")
                    import traceback
                    traceback.print_exc()
            
            # 确保必填字段
            if not title:
                title = f'{subject or "未命名"}试卷_{datetime.now().strftime("%Y%m%d")}'
            if not subject:
                subject = '其他'
            if not exam_type:
                exam_type = '其他'
            
            # 创建试卷记录
            exam_id = db.add_exam(
                title=title, subject=subject, grade=grade if grade else None, 
                exam_type=exam_type,
                exam_date=exam_date if exam_date else None, 
                total_score=total_score, 
                description=description if description else None,
                file_path=file_path, file_type=file_type, cloud_url=cloud_url,
                upload_user_id=user['id']
            )
            
            # 处理多文件上传
            if 'files' in request.files:
                files = request.files.getlist('files')
                for i, f in enumerate(files):
                    if f and f.filename and allowed_file(f.filename):
                        storage = get_storage_manager()
                        upload_result = storage.upload_file(f, folder=f'exams/{exam_id}')
                        if upload_result['success']:
                            db.add_exam_file(
                                exam_id=exam_id,
                                file_name=f.filename,
                                file_path=upload_result.get('file_path'),
                                cloud_url=upload_result.get('cloud_url'),
                                file_type=f.filename.rsplit('.', 1)[1].lower(),
                                page_number=i+1
                            )
            
            # 保存识别结果
            if recognition_result:
                # 保存题目
                for q in recognition_result.get('questions', []):
                    try:
                        db.add_question(
                            exam_id=exam_id,
                            question_number=q.get('number'),
                            question_type=q.get('type', '解答题'),
                            content=q.get('content')
                        )
                    except:
                        pass
                
                # 保存AI分析
                if recognition_result.get('ai_analysis'):
                    try:
                        db.add_ai_analysis(
                            exam_id=exam_id,
                            analysis_type='自动识别',
                            content=json.dumps(recognition_result['ai_analysis'], ensure_ascii=False, indent=2),
                            model_name='auto'
                        )
                    except:
                        pass
            
            flash(f'试卷上传成功！{subject} {grade or ""} {exam_type}', 'success')
            
            # 清除相关缓存
            invalidate_cache('dashboard:')
            invalidate_cache('exams:')
            
            return redirect(url_for('exam_detail', exam_id=exam_id))
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f'上传失败: {str(e)}', 'error')
            return redirect(url_for('upload_exam'))
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    grades = ['高一', '高二', '高三']
    exam_types = ['月考', '期中', '期末', '模拟', '高考真题', '竞赛']
    
    return render_template('upload_exam.html', subjects=subjects, grades=grades, exam_types=exam_types)

@app.route('/exams/<int:exam_id>')
@login_required
def exam_detail(exam_id):
    exam = db.get_exam_by_id(exam_id)
    if not exam:
        flash('试卷不存在', 'error')
        return redirect(url_for('exam_list'))
    
    files = db.get_exam_files(exam_id)
    questions = db.get_questions_by_exam(exam_id)
    analyses = db.get_ai_analysis(exam_id)
    
    return render_template('exam_detail.html', exam=exam, files=files, 
                          questions=questions, analyses=analyses)

@app.route('/exams/<int:exam_id>/delete', methods=['POST'])
@login_required
def delete_exam(exam_id):
    db.delete_exam(exam_id)
    flash('试卷已删除', 'success')
    
    # 清除相关缓存
    invalidate_cache('dashboard:')
    invalidate_cache('exams:')
    
    return redirect(url_for('exam_list'))

@app.route('/api/exams/<int:exam_id>/files')
@login_required
def get_exam_files_api(exam_id):
    """获取试卷文件列表API"""
    exam = db.get_exam_by_id(exam_id)
    if not exam:
        return jsonify({'success': False, 'error': '试卷不存在'})
    
    files = db.get_exam_files(exam_id)
    
    # 如果没有exam_files记录，尝试返回主文件
    if not files and exam.get('file_path'):
        files = [{
            'id': 0,
            'exam_id': exam_id,
            'file_name': exam.get('title', '试卷文件'),
            'file_path': exam['file_path'],
            'cloud_url': exam.get('cloud_url'),
            'file_type': exam.get('file_type', ''),
            'page_number': 1
        }]
    
    # 处理文件路径，生成正确的URL
    processed_files = []
    for f in files:
        file_dict = dict(f)
        # 生成可访问的URL
        if file_dict.get('cloud_url'):
            file_dict['url'] = file_dict['cloud_url']
        elif file_dict.get('file_path'):
            # 从完整路径中提取相对路径
            file_path = file_dict['file_path']
            upload_folder = app.config['UPLOAD_FOLDER']
            
            # 如果是完整路径，转换为相对路径
            if os.path.isabs(file_path):
                # 获取相对于 UPLOAD_FOLDER 的路径
                try:
                    rel_path = os.path.relpath(file_path, upload_folder)
                    file_dict['url'] = '/uploads/' + rel_path.replace('\\', '/')
                except ValueError:
                    # 如果无法计算相对路径，使用文件名
                    file_dict['url'] = '/uploads/exams/' + os.path.basename(file_path)
            else:
                # 已经是相对路径
                file_dict['url'] = '/uploads/' + file_path.replace('\\', '/')
        else:
            file_dict['url'] = ''
        
        processed_files.append(file_dict)
    
    return jsonify({
        'success': True,
        'files': processed_files
    })

@app.route('/exams/<int:exam_id>/edit', methods=['POST'])
@login_required
def edit_exam(exam_id):
    exam = db.get_exam_by_id(exam_id)
    if not exam:
        return jsonify({'success': False, 'error': '试卷不存在'})
    
    title = request.form.get('title', '').strip()
    subject = request.form.get('subject', '').strip()
    grade = request.form.get('grade', '').strip()
    exam_type = request.form.get('exam_type', '').strip()
    exam_date = request.form.get('exam_date', '').strip()
    total_score = request.form.get('total_score', type=float)
    description = request.form.get('description', '').strip()
    
    try:
        db.update_exam(
            exam_id,
            title=title if title else exam['title'],
            subject=subject if subject else exam['subject'],
            grade=grade if grade else None,
            exam_type=exam_type if exam_type else exam['exam_type'],
            exam_date=exam_date if exam_date else None,
            total_score=total_score if total_score else exam['total_score'],
            description=description if description else None
        )
        flash('试卷信息已更新', 'success')
    except Exception as e:
        flash(f'更新失败: {str(e)}', 'error')
    
    return redirect(url_for('exam_detail', exam_id=exam_id))

# ==================== AI分析 ====================
@app.route('/exams/<int:exam_id>/analyses/check')
@login_required
def check_existing_analysis(exam_id):
    """检查试卷是否已有某类型的分析"""
    analysis_type = request.args.get('type', '')
    existing = db.get_latest_ai_analysis(exam_id, analysis_type)
    if existing:
        return jsonify({
            'exists': True,
            'created_at': existing.get('created_at', ''),
            'model_name': existing.get('model_name', ''),
            'difficulty': existing.get('difficulty', ''),
            'knowledge_summary': existing.get('knowledge_summary', '')
        })
    return jsonify({'exists': False})

@app.route('/exams/<int:exam_id>/analyze', methods=['POST'])
@login_required
def analyze_exam(exam_id):
    try:
        exam = db.get_exam_by_id(exam_id)
        if not exam:
            return jsonify({'success': False, 'error': '试卷不存在'})
        
        analysis_type = request.form.get('analysis_type', '综合分析')
        questions = db.get_questions_by_exam(exam_id)
        files = db.get_exam_files(exam_id)
        user = get_current_user()
        
        print(f"开始分析试卷: {exam_id}, 类型: {analysis_type}")
        print(f"试卷标题: {exam.get('title')}")
        print(f"文件数量: {len(files)}")
        
        # 获取图片路径（包括PDF转换的图片）
        image_paths = []
        pdf_text = ""
        
        for f in files:
            file_path = f['cloud_url'] or f['file_path']
            file_type = f['file_type'].lower() if f['file_type'] else ''
            
            print(f"处理文件: {file_path}, 类型: {file_type}")
            
            if file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
                # 直接添加图片
                image_paths.append(file_path)
            elif file_type == 'pdf':
                # 处理PDF文件
                try:
                    pdf_result = convert_pdf_to_images(file_path)
                    print(f"PDF转换结果: {pdf_result['success']}")
                    if pdf_result['success']:
                        image_paths.extend(pdf_result['images'])
                        pdf_text = pdf_result.get('text', '')
                except Exception as e:
                    print(f"PDF处理失败: {e}")
                    import traceback
                    traceback.print_exc()
        
        # 如果没有文件，尝试使用主文件
        if not image_paths and exam.get('file_path'):
            file_type = exam.get('file_type', '').lower()
            print(f"使用主文件: {exam['file_path']}, 类型: {file_type}")
            
            if file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
                image_paths.append(exam['file_path'])
            elif file_type == 'pdf':
                try:
                    pdf_result = convert_pdf_to_images(exam['file_path'])
                    print(f"PDF转换结果: {pdf_result['success']}")
                    if pdf_result['success']:
                        image_paths.extend(pdf_result['images'])
                        pdf_text = pdf_result.get('text', '')
                except Exception as e:
                    print(f"PDF处理失败: {e}")
                    import traceback
                    traceback.print_exc()
        
        print(f"图片数量: {len(image_paths)}")
        
        analyzer = get_analyzer()
        print(f"AI分析器: {analyzer.provider} - {analyzer.model}")
        
        result = analyzer.analyze_exam(exam, questions, image_paths, analysis_type, exam_text=pdf_text)
        
        print(f"分析结果: {result['success']}")
        
        if result['success']:
            content = result['content']
            
            # 从AI结果中提取结构化信息
            difficulty = None
            knowledge_summary = None
            question_count = None
            
            if questions:
                question_count = len(questions)
            
            # 尝试从内容中提取难度和知识点摘要
            content_lower = content.lower()
            for diff in ['简单', '中等', '困难', '较难', '较易']:
                if diff in content:
                    difficulty = diff
                    break
            
            # 提取知识点摘要（取前200字作为摘要）
            import re as _re
            knowledge_match = _re.search(r'(?:知识点|考点|涉及知识)[：:]\s*(.{20,200})', content)
            if knowledge_match:
                knowledge_summary = knowledge_match.group(1).strip()
            
            db.add_ai_analysis(
                exam_id, analysis_type, content, result.get('model'),
                difficulty=difficulty,
                knowledge_summary=knowledge_summary,
                question_count=question_count
            )
            
            # 如果是错题提取与分析，自动保存到错题本
            saved_mistakes = []
            if analysis_type == '错题提取与分析':
                try:
                    # 解析JSON结果
                    import re
                    # 移除可能的markdown代码块标记
                    json_str = content.strip()
                    if json_str.startswith('```'):
                        json_str = json_str.split('\n', 1)[1] if '\n' in json_str else json_str[3:]
                    if json_str.endswith('```'):
                        json_str = json_str[:-3]
                    if json_str.startswith('json'):
                        json_str = json_str[4:]
                    
                    analysis_data = json.loads(json_str.strip())
                    
                    if 'mistakes' in analysis_data:
                        for mistake in analysis_data['mistakes']:
                            # 构建错题内容
                            mistake_content = f"【{mistake.get('question_type', '未知题型')}】{mistake.get('content', '')}"
                            if mistake.get('common_wrong_answer'):
                                mistake_content += f"\n\n常见错误答案：{mistake['common_wrong_answer']}"
                            
                            # 构建分析内容
                            analysis_content = f"错误分析：{mistake.get('error_analysis', '')}\n"
                            analysis_content += f"正确答案：{mistake.get('correct_answer', '')}\n"
                            analysis_content += f"解题思路：{mistake.get('solution_steps', '')}\n"
                            analysis_content += f"涉及知识点：{mistake.get('knowledge_point', '')}"
                            
                            # 保存到错题本
                            db.add_mistake(
                                user_id=user['id'],
                                question_id=None,
                                exam_id=exam_id,
                                subject=exam['subject'],
                                question_type=mistake.get('question_type', '解答题'),
                                content=mistake_content,
                                correct_answer=mistake.get('correct_answer', ''),
                                user_answer=mistake.get('common_wrong_answer', ''),
                                analysis=analysis_content
                            )
                            saved_mistakes.append(mistake.get('question_number', '?'))
                    
                except json.JSONDecodeError as e:
                    print(f"JSON解析失败: {e}")
                except Exception as e:
                    print(f"保存错题失败: {e}")
            
            response_data = {
                'success': True,
                'content': content,
                'difficulty': difficulty,
                'knowledge_summary': knowledge_summary,
                'question_count': question_count
            }
            if saved_mistakes:
                response_data['saved_mistakes'] = saved_mistakes
                response_data['message'] = f'已自动将 {len(saved_mistakes)} 道错题添加到错题本'
            
            return jsonify(response_data)
        else:
            return jsonify({'success': False, 'error': result.get('error', '分析失败')})
            
    except Exception as e:
        import traceback
        print(f"分析试卷失败: {e}")
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})

def convert_pdf_to_images(pdf_path):
    """将PDF转换为图片"""
    import fitz  # PyMuPDF
    import base64
    import tempfile
    
    result = {
        'success': False,
        'images': [],
        'text': ''
    }
    
    try:
        # 获取实际文件路径
        actual_path = pdf_path
        
        # 处理不同格式的路径
        if pdf_path.startswith('/uploads/'):
            # 转换为本地路径
            actual_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_path[9:])
        elif 'uploads' in pdf_path:
            # 包含uploads的相对路径，直接使用
            actual_path = pdf_path
        elif not os.path.isabs(pdf_path):
            # 其他相对路径
            actual_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_path)
        
        # 统一路径分隔符
        actual_path = actual_path.replace('\\\\', os.sep).replace('\\', os.sep)
        
        print(f"PDF路径转换: {pdf_path} -> {actual_path}")
        print(f"文件存在: {os.path.exists(actual_path)}")
        
        if not os.path.exists(actual_path):
            result['error'] = f"文件不存在: {actual_path}"
            return result
        
        # 打开PDF
        doc = fitz.open(actual_path)
        
        # 提取文字
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        result['text'] = full_text[:2000]  # 限制长度
        
        # 转换前几页为图片（最多5页）
        max_pages = min(5, len(doc))
        temp_dir = tempfile.mkdtemp()
        
        for page_num in range(max_pages):
            page = doc[page_num]
            # 设置缩放比例以获得合适的图片大小
            zoom = 2.0
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # 保存为临时文件
            img_path = os.path.join(temp_dir, f'page_{page_num + 1}.png')
            pix.save(img_path)
            
            # 转为base64
            with open(img_path, 'rb') as f:
                img_data = base64.b64encode(f.read()).decode()
            
            result['images'].append(f'data:image/png;base64,{img_data}')
            
            # 清理临时文件
            os.remove(img_path)
        
        os.rmdir(temp_dir)
        doc.close()
        
        result['success'] = True
        print(f"PDF转换成功: {len(result['images'])} 页")
        
    except Exception as e:
        print(f"PDF转换失败: {e}")
        result['error'] = str(e)
    
    return result

@app.route('/exams/<int:exam_id>/chat', methods=['POST'])
@login_required
def chat_about_exam(exam_id):
    exam = db.get_exam_by_id(exam_id)
    if not exam:
        return jsonify({'success': False, 'error': '试卷不存在'})
    
    message = request.form.get('message')
    context_json = request.form.get('context', '[]')
    
    try:
        context = json.loads(context_json)
    except:
        context = []
    
    # 构建上下文
    system_prompt = f"""你是一位专业的高中教师助手。当前讨论的试卷是：
标题：{exam['title']}
科目：{exam['subject']}
年级：{exam['grade']}
请用专业、耐心的态度回答学生的问题。"""
    
    messages = [{'role': 'system', 'content': system_prompt}] + context
    messages.append({'role': 'user', 'content': message})
    
    analyzer = get_analyzer()
    response = analyzer.chat(message, messages[1:])  # 去掉system prompt
    
    return jsonify({'success': True, 'response': response})

# ==================== 错题本 ====================
@app.route('/mistakes')
@login_required
def mistake_list():
    user = get_current_user()
    subject = request.args.get('subject', '')
    
    mistakes = db.get_mistakes(user['id'], subject=subject if subject else None)
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    
    return render_template('mistake_list.html', mistakes=mistakes, subjects=subjects, 
                          selected_subject=subject)

@app.route('/mistakes/add', methods=['POST'])
@login_required
def add_mistake():
    user = get_current_user()
    
    question_id = request.form.get('question_id', type=int)
    exam_id = request.form.get('exam_id', type=int)
    subject = request.form.get('subject')
    question_type = request.form.get('question_type')
    content = request.form.get('content')
    correct_answer = request.form.get('correct_answer')
    user_answer = request.form.get('user_answer')
    analysis = request.form.get('analysis')
    
    db.add_mistake(
        user_id=user['id'],
        question_id=question_id,
        exam_id=exam_id,
        subject=subject,
        question_type=question_type,
        content=content,
        correct_answer=correct_answer,
        user_answer=user_answer,
        analysis=analysis
    )
    
    flash('已添加到错题本', 'success')
    return redirect(request.referrer or url_for('mistake_list'))

@app.route('/mistakes/<int:mistake_id>/update', methods=['POST'])
@login_required
def update_mistake(mistake_id):
    mastery_level = request.form.get('mastery_level', type=int)
    db.update_mistake(mistake_id, mastery_level=mastery_level)
    return jsonify({'success': True})

@app.route('/mistakes/<int:mistake_id>/review', methods=['POST'])
@login_required
def mark_reviewed(mistake_id):
    """标记错题已复习，更新复习次数和下次复习日期"""
    user = get_current_user()
    mistake = db.get_mistake_by_id(mistake_id)
    
    if not mistake:
        return jsonify({'success': False, 'message': '错题不存在'})
    
    # 更新复习次数
    review_count = (mistake.get('review_count') or 0) + 1
    
    # 基于艾宾浩斯遗忘曲线计算下次复习日期
    from datetime import datetime, timedelta
    review_intervals = [0.014, 0.042, 0.375, 1, 2, 6, 31]  # 20分钟, 1小时, 9小时, 1天, 2天, 6天, 31天
    
    if review_count <= len(review_intervals):
        days = review_intervals[min(review_count - 1, len(review_intervals) - 1)]
    else:
        days = 31  # 超过7次复习，间隔31天
    
    next_review = datetime.now() + timedelta(days=days)
    
    # 更新数据库
    db.update_mistake(mistake_id, 
                      review_count=review_count, 
                      next_review_date=next_review.strftime('%Y-%m-%d'))
    
    return jsonify({'success': True, 'next_review': next_review.strftime('%Y-%m-%d')})

@app.route('/review-reminder')
@login_required
def review_reminder():
    """复习提醒页面"""
    user = get_current_user()
    filter_type = request.args.get('filter', 'all')
    selected_subject = request.args.get('subject', '')
    
    # 获取所有错题
    mistakes = db.get_mistakes(user['id'], subject=selected_subject if selected_subject else None)
    
    from datetime import date, datetime
    today = date.today().isoformat()
    
    # 分类错题
    review_items = []
    stats = {'urgent': 0, 'today': 0, 'upcoming': 0, 'completed': 0}
    
    for mistake in mistakes:
        next_review = mistake.get('next_review_date')
        if not next_review:
            # 没有设置复习日期，设置为今天
            next_review = today
            db.update_mistake(mistake['id'], next_review_date=today)
        
        # 确定优先级
        if next_review < today:
            priority = 'urgent'
            stats['urgent'] += 1
        elif next_review == today:
            priority = 'today'
            stats['today'] += 1
        else:
            # 检查是否在7天内
            review_date = datetime.strptime(next_review, '%Y-%m-%d').date()
            days_diff = (review_date - date.today()).days
            if days_diff <= 7:
                priority = 'upcoming'
                stats['upcoming'] += 1
            else:
                priority = 'later'
        
        # 根据筛选条件过滤
        if filter_type == 'all' or filter_type == priority:
            mistake['priority'] = priority
            mistake['next_review_date'] = next_review
            review_items.append(mistake)
    
    # 按优先级排序
    priority_order = {'urgent': 0, 'today': 1, 'upcoming': 2, 'later': 3}
    review_items.sort(key=lambda x: (priority_order.get(x['priority'], 4), x.get('next_review_date', '')))
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    
    return render_template('review_reminder.html',
                          review_items=review_items,
                          review_stats=stats,
                          subjects=subjects,
                          selected_subject=selected_subject,
                          filter=filter_type)

# ==================== 学习计划 ====================
@app.route('/study-plan')
@login_required
def study_plan():
    """学习计划页面"""
    user = get_current_user()
    
    # 获取计划列表
    plans = db.get_study_plans(user['id'])
    
    # 获取统计信息
    stats = db.get_plan_stats(user['id'])
    
    # 获取今日任务
    today_tasks = []
    for plan in plans:
        if plan['status'] == '进行中':
            for task in plan.get('tasks', []):
                if not task['completed']:
                    task['plan_id'] = plan['id']
                    today_tasks.append(task)
    
    # 学习建议
    suggestions = [
        '每天保持2-3小时的专注学习时间',
        '错题复习比做新题更重要',
        '利用碎片时间复习知识点',
        '定期总结学习成果和不足'
    ]
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    
    # 生成日历数据
    from datetime import date, timedelta
    today = date.today()
    first_day = today.replace(day=1)
    weekday = first_day.weekday()
    
    calendar_days = []
    # 添加上个月的日期
    for i in range(weekday):
        day = first_day - timedelta(days=weekday - i)
        calendar_days.append({
            'day': day.day,
            'date': day.isoformat(),
            'is_today': False,
            'has_plan': False
        })
    
    # 添加本月的日期
    for i in range(31):
        try:
            day = first_day + timedelta(days=i)
            if day.month != today.month:
                break
            calendar_days.append({
                'day': day.day,
                'date': day.isoformat(),
                'is_today': day == today,
                'has_plan': any(p['start_date'] <= day.isoformat() <= p['end_date'] for p in plans)
            })
        except:
            break
    
    return render_template('study_plan.html',
                          plans=plans,
                          stats=stats,
                          today_tasks=today_tasks,
                          suggestions=suggestions,
                          subjects=subjects,
                          calendar_days=calendar_days,
                          current_month=today.strftime('%Y年%m月'))

@app.route('/plans/add', methods=['POST'])
@login_required
def add_study_plan():
    """添加学习计划"""
    user = get_current_user()
    
    title = request.form.get('title')
    subject = request.form.get('subject')
    plan_type = request.form.get('plan_type', '日常学习')
    start_date = request.form.get('start_date')
    end_date = request.form.get('end_date')
    daily_goal = request.form.get('daily_goal')
    description = request.form.get('description')
    tasks_text = request.form.get('tasks', '')
    
    # 创建计划
    plan_id = db.add_study_plan(
        user_id=user['id'],
        title=title,
        subject=subject,
        start_date=start_date,
        end_date=end_date,
        plan_type=plan_type,
        description=description,
        daily_goal=daily_goal
    )
    
    # 添加任务
    if tasks_text.strip():
        tasks = [t.strip() for t in tasks_text.split('\n') if t.strip()]
        for i, task_content in enumerate(tasks):
            db.add_plan_task(plan_id, task_content, sort_order=i)
    
    flash('学习计划创建成功！', 'success')
    return redirect(url_for('study_plan'))

@app.route('/plans/<int:plan_id>/complete', methods=['POST'])
@login_required
def complete_plan(plan_id):
    """完成学习计划"""
    db.update_study_plan(plan_id, status='已完成', progress=100)
    return jsonify({'success': True})

@app.route('/plans/<int:plan_id>/delete', methods=['POST'])
@login_required
def delete_plan(plan_id):
    """删除学习计划"""
    db.delete_study_plan(plan_id)
    return jsonify({'success': True})

@app.route('/plans/<int:plan_id>/tasks/<int:task_id>/toggle', methods=['POST'])
@login_required
def toggle_task(plan_id, task_id):
    """切换任务完成状态"""
    data = request.get_json()
    completed = data.get('completed', False)
    
    db.update_plan_task(task_id, completed=completed)
    
    # 重新计算进度
    tasks = db.get_plan_tasks(plan_id)
    if tasks:
        completed_count = sum(1 for t in tasks if t['completed'])
        progress = round(completed_count / len(tasks) * 100, 1)
        db.update_study_plan(plan_id, progress=progress)
    else:
        progress = 0
    
    return jsonify({'success': True, 'progress': progress})

@app.route('/plans/<int:plan_id>/edit')
@login_required
def edit_plan(plan_id):
    """编辑学习计划（可以在这里实现编辑功能）"""
    # 简化处理：重定向回计划页面
    flash('编辑功能开发中', 'info')
    return redirect(url_for('study_plan'))

# ==================== 分析管理 ====================
@app.route('/analyses')
@login_required
def analysis_list():
    exam_id = request.args.get('exam_id', type=int)
    analysis_type = request.args.get('analysis_type', '')
    
    analyses = db.get_all_ai_analyses(
        exam_id=exam_id,
        analysis_type=analysis_type if analysis_type else None
    )
    
    analysis_types = ['综合分析', '知识点总结', '学习建议', '错题提取与分析', '错题分析', '自动识别']
    
    return render_template('analysis_list.html', analyses=analyses,
                          analysis_types=analysis_types,
                          selected_type=analysis_type,
                          selected_exam_id=exam_id)

@app.route('/analyses/<int:analysis_id>/delete', methods=['POST'])
@login_required
def delete_analysis(analysis_id):
    db.delete_ai_analysis(analysis_id)
    flash('分析记录已删除', 'success')
    return redirect(url_for('analysis_list'))

# ==================== 学习统计 ====================
@app.route('/stats')
@login_required
def study_stats():
    user = get_current_user()
    subject = request.args.get('subject', '')
    
    stats = db.get_study_stats(user['id'], subject=subject if subject else None)
    subject_stats = db.get_subject_stats(user['id'])
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    
    return render_template('study_stats.html', stats=stats, subject_stats=subject_stats,
                          subjects=subjects, selected_subject=subject)

@app.route('/stats/report')
@login_required
def study_report():
    """生成学习报告"""
    user = get_current_user()
    from datetime import datetime
    
    # 获取统计数据
    all_exams = db.get_all_exams(limit=10000)
    mistakes = db.get_mistakes(user['id'])
    
    # 计算正确率
    total_questions = 0
    correct_questions = 0
    for exam in all_exams:
        questions = db.get_questions(exam['id'])
        total_questions += len(questions)
        correct_questions += sum(1 for q in questions if q.get('user_score') and q['user_score'] > 0)
    
    accuracy_rate = round(correct_questions / total_questions * 100, 1) if total_questions > 0 else 0
    
    stats = {
        'total_exams': len(all_exams),
        'total_questions': total_questions,
        'total_mistakes': len(mistakes),
        'accuracy_rate': accuracy_rate,
    }
    
    # 各科成绩统计
    subject_stats = db.get_subject_stats(user['id'])
    
    # 错题统计
    mastered_count = sum(1 for m in mistakes if m.get('mastery_level', 0) >= 4)
    review_count = sum(1 for m in mistakes if m.get('next_review_date') and m['next_review_date'] <= datetime.now().strftime('%Y-%m-%d'))
    
    # 各科错题分布
    mistake_subjects = {}
    for m in mistakes:
        subject = m.get('subject', '未知')
        mistake_subjects[subject] = mistake_subjects.get(subject, 0) + 1
    mistake_subject_list = [{'subject': k, 'count': v} for k, v in mistake_subjects.items()]
    
    # 学习计划统计
    plan_stats = db.get_plan_stats(user['id'])
    
    # 学习建议
    suggestions = []
    if accuracy_rate < 60:
        suggestions.append('正确率较低，建议加强基础知识复习')
    if len(mistakes) > 50:
        suggestions.append('错题较多，建议定期复习错题本')
    if plan_stats.get('completion_rate', 0) < 50:
        suggestions.append('学习计划完成率较低，建议制定更合理的计划')
    if not suggestions:
        suggestions.append('继续保持良好的学习习惯')
        suggestions.append('适当增加练习难度，挑战更高目标')
    
    return render_template('study_report.html',
                          user=user,
                          stats=stats,
                          subject_stats=subject_stats,
                          mastered_count=mastered_count,
                          review_count=review_count,
                          mistake_subjects=mistake_subject_list,
                          plan_stats=plan_stats,
                          suggestions=suggestions,
                          report_date=datetime.now().strftime('%Y年%m月%d日'),
                          generated_at=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

@app.route('/stats/add', methods=['POST'])
@login_required
def add_study_stat():
    user = get_current_user()
    
    subject = request.form.get('subject')
    exam_date = request.form.get('exam_date')
    score = request.form.get('score', type=float)
    total_score = request.form.get('total_score', type=float)
    rank = request.form.get('rank', type=int)
    class_avg = request.form.get('class_avg', type=float)
    
    db.add_study_stat(user['id'], subject, exam_date, score, total_score, rank, class_avg)
    
    flash('成绩已记录', 'success')
    return redirect(url_for('study_stats'))

# ==================== 专项练习 ====================
@app.route('/practice')
@login_required
def practice_list():
    """练习列表"""
    user = get_current_user()
    subject = request.args.get('subject', '')
    status = request.args.get('status', '')
    
    sessions = db.get_user_practice_sessions(
        user['id'],
        subject=subject if subject else None,
        status=status if status else None
    )
    
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    
    return render_template('practice_list.html', sessions=sessions,
                          subjects=subjects, selected_subject=subject,
                          selected_status=status)

@app.route('/practice/start/<int:mistake_id>', methods=['POST'])
@login_required
def start_practice(mistake_id):
    """从错题开始练习"""
    user = get_current_user()
    
    # 获取错题信息
    mistakes = db.get_mistakes(user['id'])
    mistake = None
    for m in mistakes:
        if m['id'] == mistake_id:
            mistake = m
            break
    
    if not mistake:
        flash('错题不存在', 'error')
        return redirect(url_for('mistake_list'))
    
    # 创建练习会话
    session_id = db.create_practice_session(
        user_id=user['id'],
        mistake_id=mistake_id,
        subject=mistake['subject'],
        practice_type='错题巩固',
        title=f"针对错题的专项练习"
    )
    
    # 使用AI生成练习题
    analyzer = get_analyzer()
    result = analyzer.generate_practice_questions(mistake, num_questions=3)
    
    question_number = 0
    if result and 'questions' in result:
        for q in result['questions']:
            question_number += 1
            db.add_practice_question(
                session_id=session_id,
                question_number=question_number,
                question_type=q.get('question_type', '解答题'),
                content=q.get('content', ''),
                options=q.get('options'),
                correct_answer=q.get('correct_answer', ''),
                analysis=q.get('analysis', '')
            )
    
    # 搜索匹配的高考真题（按科目和知识点）
    knowledge_point = mistake.get('analysis', '') or mistake.get('content', '')
    # 从错题内容中提取关键词作为知识点
    kp_keywords = ''
    if '知识点' in knowledge_point:
        import re as _re
        m = _re.search(r'知识点[：:]\s*(.+?)(?:\n|$)', knowledge_point)
        if m:
            kp_keywords = m.group(1).strip()
    if not kp_keywords:
        kp_keywords = mistake.get('content', '')[:50]
    
    gaokao_matches = db.search_gaokao_questions(
        subject=mistake.get('subject'),
        knowledge_point=kp_keywords,
        limit=2
    )
    
    for gq in gaokao_matches:
        question_number += 1
        db.add_practice_question(
            session_id=session_id,
            question_number=question_number,
            question_type=gq.get('question_type', '解答题'),
            content=f"【高考真题·{gq.get('year', '')}年{gq.get('region', '')}】\n{gq.get('content', '')}",
            options=gq.get('options'),
            correct_answer=gq.get('correct_answer', ''),
            analysis=gq.get('analysis', '')
        )
    
    if question_number > 0:
        return redirect(url_for('do_practice', session_id=session_id))
    else:
        flash('生成练习题失败，请重试', 'error')
        return redirect(url_for('mistake_list'))

@app.route('/practice/<int:session_id>')
@login_required
def do_practice(session_id):
    """做练习"""
    user = get_current_user()
    session = db.get_practice_session(session_id)
    
    if not session or session['user_id'] != user['id']:
        flash('练习不存在', 'error')
        return redirect(url_for('practice_list'))
    
    questions = db.get_practice_questions(session_id)
    
    return render_template('do_practice.html', practice_session=session, questions=questions)

@app.route('/practice/<int:session_id>/submit', methods=['POST'])
@login_required
def submit_practice(session_id):
    """提交练习"""
    user = get_current_user()
    session = db.get_practice_session(session_id)
    
    if not session or session['user_id'] != user['id']:
        return jsonify({'success': False, 'error': '练习不存在'})
    
    questions = db.get_practice_questions(session_id)
    correct_count = 0
    
    for q in questions:
        user_answer = request.form.get(f'answer_{q["id"]}', '').strip()
        if not user_answer:
            continue
        
        # 简单的答案比对（可以根据需要改进）
        is_correct = 1 if user_answer.lower().strip() == q['correct_answer'].lower().strip() else 0
        
        db.update_practice_answer(q['id'], user_answer, is_correct)
        
        if is_correct:
            correct_count += 1
    
    # 完成练习
    score = db.complete_practice_session(session_id, correct_count, len(questions))
    
    # 如果有错题关联，更新错题的复习次数
    if session['mistake_id']:
        db.update_mistake(session['mistake_id'], review_count=db.get_mistakes(user['id'])[0].get('review_count', 0) + 1)
    
    flash(f'练习完成！得分：{score:.1f}%（{correct_count}/{len(questions)}）', 'success')
    return redirect(url_for('practice_result', session_id=session_id))

@app.route('/practice/<int:session_id>/result')
@login_required
def practice_result(session_id):
    """练习结果"""
    user = get_current_user()
    session = db.get_practice_session(session_id)
    
    if not session or session['user_id'] != user['id']:
        flash('练习不存在', 'error')
        return redirect(url_for('practice_list'))
    
    questions = db.get_practice_questions(session_id)
    
    return render_template('practice_result.html', practice_session=session, questions=questions)

@app.route('/practice/<int:session_id>/correction', methods=['POST'])
@login_required
def submit_correction(session_id):
    """提交改错"""
    user = get_current_user()
    session = db.get_practice_session(session_id)
    
    if not session or session['user_id'] != user['id']:
        return jsonify({'success': False, 'error': '练习不存在'})
    
    question_id = request.form.get('question_id', type=int)
    user_correction = request.form.get('correction', '').strip()
    
    question = None
    for q in db.get_practice_questions(session_id):
        if q['id'] == question_id:
            question = q
            break
    
    if not question:
        return jsonify({'success': False, 'error': '题目不存在'})
    
    # 判断改错是否正确
    is_correct = 1 if user_correction.lower().strip() == question['correct_answer'].lower().strip() else 0
    
    # 添加改错记录
    db.add_correction_record(
        user_id=user['id'],
        mistake_id=session['mistake_id'],
        practice_question_id=question_id,
        session_id=session_id,
        original_mistake=question['content'],
        user_correction=user_correction,
        is_correct=is_correct
    )
    
    return jsonify({
        'success': True,
        'is_correct': bool(is_correct),
        'correct_answer': question['correct_answer'],
        'analysis': question['analysis']
    })

@app.route('/corrections')
@login_required
def correction_list():
    """改错管理"""
    user = get_current_user()
    mistake_id = request.args.get('mistake_id', type=int)
    
    corrections = db.get_correction_records(user['id'], mistake_id=mistake_id)
    
    return render_template('correction_list.html', corrections=corrections,
                          selected_mistake_id=mistake_id)

# ==================== 系统设置 ====================
@app.route('/settings')
@login_required
def settings():
    user = get_current_user()
    if user['role'] != 'admin':
        flash('权限不足', 'error')
        return redirect(url_for('dashboard'))
    
    configs = db.get_all_configs()
    return render_template('settings.html', configs=configs)

@app.route('/settings/update', methods=['POST'])
@login_required
def update_settings():
    user = get_current_user()
    if user['role'] != 'admin':
        return jsonify({'success': False, 'error': '权限不足'})
    
    key = request.form.get('key')
    value = request.form.get('value')
    description = request.form.get('description')
    
    db.set_config(key, value, description)
    
    return jsonify({'success': True})

@app.route('/settings/test-ai', methods=['POST'])
@login_required
def test_ai():
    user = get_current_user()
    if user['role'] != 'admin':
        return jsonify({'success': False, 'error': '权限不足'})
    
    analyzer = get_analyzer()
    result = analyzer.chat('你好，请回复"连接成功"四个字。')
    
    if '连接成功' in result or '成功' in result:
        return jsonify({'success': True, 'message': 'AI连接测试成功！'})
    else:
        return jsonify({'success': False, 'message': f'AI响应: {result}'})

# ==================== 用户管理 ====================
@app.route('/users')
@login_required
def user_list():
    user = get_current_user()
    if user['role'] != 'admin':
        flash('权限不足', 'error')
        return redirect(url_for('dashboard'))
    
    users = db.get_all_users()
    return render_template('user_list.html', users=users)

@app.route('/users/add', methods=['POST'])
@login_required
def add_user():
    user = get_current_user()
    if user['role'] != 'admin':
        return jsonify({'success': False, 'error': '权限不足'})
    
    username = request.form.get('username')
    password = request.form.get('password')
    role = request.form.get('role', 'student')
    display_name = request.form.get('display_name')
    grade = request.form.get('grade')
    class_name = request.form.get('class_name')
    
    ok, msg = db.add_user(username, password, role, display_name, grade, class_name)
    
    if ok:
        flash('用户添加成功', 'success')
    else:
        flash(f'添加失败: {msg}', 'error')
    
    return redirect(url_for('user_list'))

# ==================== 高考真题 ====================
@app.route('/gaokao')
@login_required
def gaokao_list():
    """高考真题库（分类浏览）"""
    subject = request.args.get('subject', '')
    category = request.args.get('category', '')
    question_type = request.args.get('question_type', '')
    difficulty = request.args.get('difficulty', '')
    year = request.args.get('year', '', type=str)
    keyword = request.args.get('keyword', '')
    
    questions = db.get_all_gaokao_questions(
        subject=subject if subject else None,
        year=int(year) if year else None,
        category=category if category else None,
        question_type=question_type if question_type else None,
        difficulty=difficulty if difficulty else None
    )
    
    # 关键词搜索（在内存中过滤）
    if keyword:
        keyword_lower = keyword.lower()
        questions = [q for q in questions if 
                    keyword_lower in (q.get('content', '') or '').lower() or
                    keyword_lower in (q.get('knowledge_point', '') or '').lower() or
                    keyword_lower in (q.get('category', '') or '').lower() or
                    keyword_lower in (q.get('analysis', '') or '').lower()]
    
    stats = db.get_gaokao_stats()
    category_tree = db.get_gaokao_category_tree()
    subjects = ['语文', '数学', '英语', '物理', '化学', '生物', '历史', '地理', '政治']
    
    return render_template('gaokao_list.html', 
                          questions=questions, stats=stats, 
                          category_tree=category_tree,
                          subjects=subjects,
                          selected_subject=subject,
                          selected_category=category,
                          selected_question_type=question_type,
                          selected_difficulty=difficulty,
                          selected_year=year,
                          keyword=keyword)

@app.route('/gaokao/add', methods=['POST'])
@login_required
def gaokao_add():
    """添加高考真题"""
    try:
        question_id = db.add_gaokao_question(
            subject=request.form.get('subject', '数学'),
            category=request.form.get('category', ''),
            year=request.form.get('year', type=int),
            region=request.form.get('region', ''),
            question_number=request.form.get('question_number', type=int),
            question_type=request.form.get('question_type', '选择题'),
            content=request.form.get('content', ''),
            options=request.form.get('options', ''),
            correct_answer=request.form.get('correct_answer', ''),
            analysis=request.form.get('analysis', ''),
            knowledge_point=request.form.get('knowledge_point', ''),
            difficulty=request.form.get('difficulty', '中等'),
            score=request.form.get('score', type=float),
            source=request.form.get('source', '')
        )
        flash(f'真题添加成功 (ID: {question_id})', 'success')
    except Exception as e:
        flash(f'添加失败: {str(e)}', 'error')
    return redirect(url_for('gaokao_list'))

@app.route('/gaokao/<int:question_id>/edit', methods=['POST'])
@login_required
def gaokao_edit(question_id):
    """编辑高考真题"""
    try:
        kw = {}
        for field in ['subject', 'category', 'region', 'question_type', 'content', 
                       'options', 'correct_answer', 'analysis', 'knowledge_point', 'difficulty', 'source']:
            val = request.form.get(field)
            if val is not None:
                kw[field] = val
        if request.form.get('year'):
            kw['year'] = request.form.get('year', type=int)
        if request.form.get('question_number'):
            kw['question_number'] = request.form.get('question_number', type=int)
        if request.form.get('score'):
            kw['score'] = request.form.get('score', type=float)
        
        if kw:
            db.update_gaokao_question(question_id, **kw)
        flash('真题已更新', 'success')
    except Exception as e:
        flash(f'更新失败: {str(e)}', 'error')
    return redirect(url_for('gaokao_list'))

@app.route('/gaokao/<int:question_id>/delete', methods=['POST'])
@login_required
def gaokao_delete(question_id):
    """删除高考真题"""
    db.delete_gaokao_question(question_id)
    flash('真题已删除', 'success')
    return redirect(url_for('gaokao_list'))

@app.route('/gaokao/batch-classify', methods=['POST'])
@login_required
def gaokao_batch_classify():
    """批量AI自动分类"""
    question_ids = request.form.getlist('question_ids')
    if not question_ids:
        # 分类所有未分类的
        questions = db.get_all_gaokao_questions(limit=500)
        questions = [q for q in questions if not q.get('category')]
        question_ids = [str(q['id']) for q in questions]
    
    if not question_ids:
        flash('没有需要分类的题目', 'info')
        return redirect(url_for('gaokao_list'))
    
    analyzer = get_analyzer()
    # 批量处理，每批20题
    batch_size = 20
    total_classified = 0
    
    for i in range(0, len(question_ids), batch_size):
        batch_ids = question_ids[i:i+batch_size]
        batch_questions = []
        for qid in batch_ids:
            q = db.get_gaokao_question_by_id(int(qid))
            if q and not q.get('category'):
                batch_questions.append(q)
        
        if not batch_questions:
            continue
        
        # 构建AI分类请求
        q_list = []
        for q in batch_questions:
            q_list.append({
                'id': q['id'],
                'subject': q['subject'],
                'question_type': q['question_type'],
                'content': q['content'][:200],
                'knowledge_point': q.get('knowledge_point', '')
            })
        
        prompt = f"""请为以下高考题目进行智能分类，包括知识模块、具体知识点和难度评估。

题目列表：
{json.dumps(q_list, ensure_ascii=False, indent=2)}

分类体系（根据科目选择合适的分类）：

【数学】
一级分类：集合与逻辑、函数与导数、三角函数、数列、立体几何、解析几何、概率与统计、不等式、向量、复数
示例知识点：集合运算、函数性质、三角恒等变换、等差数列、空间向量、椭圆方程、古典概型

【英语】
一级分类：听力、阅读理解、完形填空、语法填空、短文改错、书面表达、七选五
示例知识点：细节理解、推理判断、动词时态、非谓语、从句

【物理】
一级分类：力学、电磁学、热学、光学、原子物理、实验题
示例知识点：牛顿定律、动能定理、电场强度、电磁感应、气体性质

【化学】
一级分类：有机化学、无机化学、化学反应原理、化学实验、物质结构
示例知识点：有机推断、离子反应、化学平衡、实验设计、元素周期律

【生物】
一级分类：细胞生物学、遗传学、生态学、动物生理、植物生理、生物技术
示例知识点：细胞结构、遗传规律、种群群落、神经调节、光合作用

【语文】
一级分类：现代文阅读、古诗文阅读、语言文字运用、写作、文学类文本、实用类文本
示例知识点：论述类文本、文言文阅读、成语运用、材料作文

【历史】
一级分类：中国古代史、中国近代史、中国现代史、世界古代史、世界近代史、世界现代史
示例知识点：秦汉制度、辛亥革命、工业革命、冷战格局

【地理】
一级分类：自然地理、人文地理、区域地理、地理信息技术
示例知识点：大气环流、城市化、中国地理、遥感技术

【政治】
一级分类：经济生活、政治生活、文化生活、生活与哲学
示例知识点：消费心理、政府职能、文化传承、唯物辩证法

请返回JSON格式：
{{
    "classifications": [
        {{
            "id": 题目ID,
            "category": "一级分类名称（知识模块）",
            "knowledge_points": ["具体知识点1", "具体知识点2"],
            "difficulty": "简单/中等/困难"
        }}
    ]
}}

要求：
1. category必须是上述一级分类之一
2. knowledge_points要具体明确
3. difficulty要准确评估
只返回JSON。"""

        result = analyzer.chat(prompt)
        if not result or not result.get('success'):
            continue
        
        content = result.get('content', '')
        try:
            json_str = content.strip()
            if json_str.startswith('```'):
                json_str = json_str.split('\n', 1)[1] if '\n' in json_str else json_str[3:]
            if json_str.endswith('```'):
                json_str = json_str[:-3]
            if json_str.startswith('json'):
                json_str = json_str[4:]
            
            data = json.loads(json_str.strip())
            for item in data.get('classifications', []):
                qid = item.get('id')
                cat = item.get('category', '')
                knowledge_points = item.get('knowledge_points', [])
                difficulty = item.get('difficulty', '中等')
                
                if qid and cat:
                    update_data = {'category': cat}
                    if knowledge_points:
                        update_data['knowledge_point'] = ','.join(knowledge_points) if isinstance(knowledge_points, list) else knowledge_points
                    if difficulty in ['简单', '中等', '困难']:
                        update_data['difficulty'] = difficulty
                    db.update_gaokao_question(qid, **update_data)
                    total_classified += 1
        except Exception as e:
            print(f"分类解析失败: {e}")
    
    flash(f'已自动分类 {total_classified} 道题目', 'success')
    return redirect(url_for('gaokao_list'))

# 导入任务进度存储
import_tasks = {}

@app.route('/gaokao/import', methods=['GET', 'POST'])
@login_required
def gaokao_import():
    """上传文件导入高考真题（异步处理）"""
    if request.method == 'POST':
        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({'success': False, 'error': '请选择文件'})
        
        # 保存上传的文件到临时目录
        import tempfile
        import uuid
        task_id = str(uuid.uuid4())[:8]
        tmp_dir = tempfile.mkdtemp(prefix='gaokao_')
        
        saved_files = []
        for file in files:
            if file and file.filename:
                safe_name = file.filename.replace('..', '_').replace('/', '_').replace('\\', '_')
                file_path = os.path.join(tmp_dir, safe_name)
                file.save(file_path)
                saved_files.append({'name': file.filename, 'path': file_path})
        
        if not saved_files:
            return jsonify({'success': False, 'error': '没有有效文件'})
        
        # 初始化任务状态
        import_tasks[task_id] = {
            'status': 'processing',
            'total': len(saved_files),
            'current': 0,
            'current_file': '',
            'imported': 0,
            'errors': [],
            'done': False
        }
        
        # 后台线程处理
        import threading
        thread = threading.Thread(target=_process_gaokao_import, args=(task_id, saved_files, tmp_dir))
        thread.daemon = True
        thread.start()
        
        return jsonify({'success': True, 'task_id': task_id})
    
    # GET: 显示上传页面
    return render_template('gaokao_import.html')


@app.route('/gaokao/import/progress/<task_id>')
@login_required
def gaokao_import_progress(task_id):
    """查询导入进度"""
    task = import_tasks.get(task_id)
    if not task:
        return jsonify({'success': False, 'error': '任务不存在'})
    return jsonify({'success': True, **task})


def _process_gaokao_import(task_id, saved_files, tmp_dir):
    """后台处理导入任务"""
    analyzer = get_analyzer()
    task = import_tasks[task_id]
    total_imported = 0
    
    for idx, finfo in enumerate(saved_files):
        task['current'] = idx + 1
        task['current_file'] = finfo['name']
        filename = finfo['name'].lower()
        
        try:
            text_content = ''
            
            if filename.endswith(('.doc', '.docx')):
                text_content = _extract_docx_text(finfo['path'])
            elif filename.endswith('.pdf'):
                text_content = _extract_pdf_text(finfo['path'])
            elif filename.endswith('.txt'):
                with open(finfo['path'], 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
            else:
                task['errors'].append(f'{finfo["name"]}: 不支持的文件格式')
                continue
            
            if not text_content or len(text_content.strip()) < 50:
                if filename.endswith('.doc') and not filename.endswith('.docx'):
                    task['errors'].append(f'{finfo["name"]}: .doc格式需要安装pywin32库，或转换为.docx格式')
                else:
                    task['errors'].append(f'{finfo["name"]}: 文件内容过少或无法提取文字')
                continue
            
            text_content = text_content[:8000]
            
            # 使用新的分类系统生成更详细的提示
            prompt = f"""请分析以下高考真题试卷内容，自动识别试卷信息并提取所有题目，并进行智能分类。

试卷内容：
{text_content}

请严格按照以下JSON格式返回：
{{
    "exam_info": {{
        "subject": "科目（数学/英语/物理/化学/生物/语文/历史/地理/政治）",
        "year": 2024,
        "region": "地区（如：全国甲卷、北京卷、新高考I卷等）",
        "total_score": 150,
        "exam_name": "完整试卷名称"
    }},
    "questions": [
        {{
            "question_number": 1,
            "question_type": "选择题/填空题/解答题",
            "content": "题目内容",
            "options": "A. xxx\\nB. xxx\\nC. xxx\\nD. xxx（仅选择题填写）",
            "correct_answer": "正确答案",
            "analysis": "解题分析",
            "knowledge_point": "具体知识点1,具体知识点2",
            "category": "知识模块分类",
            "difficulty": "简单/中等/困难",
            "score": 5
        }}
    ]
}}

分类要求：
1. category（知识模块）必须是该科目下的标准分类，参考以下分类体系：
- 数学：集合与逻辑、函数与导数、三角函数、数列、立体几何、解析几何、概率与统计、不等式、向量、复数
- 语文：现代文阅读、古诗文阅读、语言文字运用、写作、文学类文本、实用类文本
- 英语：听力、阅读理解、完形填空、语法填空、短文改错、书面表达、七选五
- 物理：力学、电磁学、热学、光学、原子物理、实验题
- 化学：有机化学、无机化学、化学反应原理、化学实验、物质结构
- 生物：细胞生物学、遗传学、生态学、动物生理、植物生理、生物技术
- 历史：中国古代史、中国近代史、中国现代史、世界古代史、世界近代史、世界现代史
- 地理：自然地理、人文地理、区域地理、地理信息技术
- 政治：经济生活、政治生活、文化生活、生活与哲学

2. knowledge_point（知识点）要具体，如"三角恒等变换"、"牛顿第二定律"、"有机推断"等
3. difficulty（难度）要准确评估：简单/中等/困难
4. question_type要准确识别：选择题/填空题/解答题

只返回JSON。"""

            content = None
            for attempt in range(3):  # 最多重试3次
                try:
                    result = analyzer.chat(prompt)
                    # chat()返回字符串，检查是否有效
                    if result and isinstance(result, str) and not result.startswith('错误') and '暂时不可用' not in result:
                        content = result
                        break
                    print(f"[重试] 第{attempt+1}次AI返回无效: {result[:100]}")
                except Exception as e:
                    print(f"[重试] 第{attempt+1}次AI调用异常: {e}")
                if attempt < 2:
                    import time
                    time.sleep(2)  # 等待2秒后重试
            
            if not content:
                task['errors'].append(f'{finfo["name"]}: AI分析失败（已重试3次）')
                continue
            print(f"[DEBUG] AI返回内容长度: {len(content)}")
            print(f"[DEBUG] AI返回内容前1000字: {content[:1000]}")
            print(f"[DEBUG] AI返回内容后500字: {content[-500:]}")
            json_str = content.strip()
            
            # 多种方式尝试提取JSON
            import re
            import json
            
            # 方法1: 直接解析
            data = None
            try:
                data = json.loads(json_str, strict=False)
                print(f"[DEBUG] 方法1直接解析成功")
            except Exception as e1:
                print(f"[DEBUG] 方法1失败: {e1}")
                pass
            
            # 方法2: 提取```json...```代码块
            if not data:
                code_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', json_str)
                if code_match:
                    try:
                        data = json.loads(code_match.group(1), strict=False)
                    except:
                        pass
            
            # 方法3: 提取最外层的{}
            if not data:
                brace_match = re.search(r'\{[\s\S]*\}', json_str)
                if brace_match:
                    try:
                        data = json.loads(brace_match.group(), strict=False)
                        print(f"[DEBUG] 方法3提取大括号成功")
                    except Exception as e3:
                        print(f"[DEBUG] 方法3失败: {e3}")
                        pass
            
            # 方法4: 尝试修复常见JSON错误（尾部逗号、单引号等）
            if not data:
                fixed = json_str
                # 清理Unicode控制字符（\u0000-\u001f）
                fixed = re.sub(r'[\x00-\x1f]', '', fixed)
                # 修复尾部逗号
                fixed = re.sub(r",\s*}", "}", fixed)
                fixed = re.sub(r",\s*]", "]", fixed)
                fixed = fixed.replace("'", '"')
                try:
                    data = json.loads(fixed, strict=False)
                    print(f"[DEBUG] 方法4修复后解析成功")
                except Exception as e4:
                    print(f"[DEBUG] 方法4失败: {e4}")
                    pass
            
            # 方法5: 尝试修复不完整的JSON（截断情况）
            if not data:
                fixed = json_str
                # 清理控制字符
                fixed = re.sub(r'[\x00-\x1f]', '', fixed)
                # 去掉尾部不完整的对象/数组
                fixed = re.sub(r',\s*\{[^}]*$', '', fixed)
                fixed = re.sub(r',\s*"[^"]*$', '', fixed)
                # 补全缺失的括号
                open_braces = fixed.count('{') - fixed.count('}')
                open_brackets = fixed.count('[') - fixed.count(']')
                if open_braces > 0 or open_brackets > 0:
                    # 去掉尾部逗号
                    fixed = fixed.rstrip().rstrip(',')
                    fixed += ']' * open_brackets + '}' * open_braces
                try:
                    data = json.loads(fixed, strict=False)
                    print(f"[DEBUG] 方法5修复截断JSON成功")
                except Exception as e5:
                    print(f"[DEBUG] 方法5失败: {e5}")
                    pass
            
            if not data:
                print(f"[DEBUG] JSON解析失败，原始内容: {json_str[:300]}")
                print(f"[DEBUG] JSON结尾内容: {json_str[-300:]}")
                task['errors'].append(f'{finfo["name"]}: AI返回格式错误，无法解析JSON')
                continue
            
            # 检查解析结果是否为字典
            if not isinstance(data, dict):
                print(f"[DEBUG] 解析结果不是字典: {type(data)}")
                task['errors'].append(f'{finfo["name"]}: AI返回的不是JSON字典格式')
                continue
            
            exam_info = data.get('exam_info', {})
            questions = data.get('questions', [])
            print(f"[DEBUG] 解析结果: exam_info={exam_info}, questions数量={len(questions)}")
            
            subject = exam_info.get('subject', '未知')
            year = exam_info.get('year')
            region = exam_info.get('region', '')
            source = exam_info.get('exam_name', finfo['name'])
            
            file_count = 0
            for q in questions:
                db.add_gaokao_question(
                    subject=subject,
                    category=q.get('category', ''),
                    year=year,
                    region=region,
                    question_number=q.get('question_number'),
                    question_type=q.get('question_type', '解答题'),
                    content=q.get('content', ''),
                    options=q.get('options', ''),
                    correct_answer=q.get('correct_answer', ''),
                    analysis=q.get('analysis', ''),
                    knowledge_point=q.get('knowledge_point', ''),
                    difficulty=q.get('difficulty', '中等'),
                    score=q.get('score'),
                    source=source
                )
                file_count += 1
            
            total_imported += file_count
            task['imported'] = total_imported
            print(f"导入成功: {finfo['name']} -> {subject} {year}年 {region} {file_count}题")
            
        except json.JSONDecodeError:
            task['errors'].append(f'{finfo["name"]}: AI返回格式错误')
        except Exception as e:
            task['errors'].append(f'{finfo["name"]}: {str(e)}')
            print(f"导入失败: {e}")
    
    # 清理临时文件
    import shutil
    try:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    except:
        pass
    
    task['status'] = 'done'
    task['done'] = True
    task['current_file'] = ''
    print(f"导入任务完成: {task_id}, 共导入 {total_imported} 题")


def _extract_docx_text(file_path):
    """从doc/docx文件提取文字"""
    import os
    ext = os.path.splitext(file_path)[1].lower()
    
    # .doc格式使用win32com（Windows COM接口）
    if ext == '.doc':
        try:
            import win32com.client
            import pythoncom
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            pythoncom.CoUninitialize()
            return text
        except Exception as e:
            print(f".doc提取失败(win32com): {e}")
            # 回退：尝试用antiword
            try:
                import subprocess
                result = subprocess.run(['antiword', file_path], capture_output=True, text=True, encoding='utf-8')
                if result.returncode == 0:
                    return result.stdout
            except:
                pass
            return ''
    
    # .docx格式使用python-docx
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # 也提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_text = '  '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)
    except Exception as e:
        print(f"docx提取失败: {e}")
        return ''


def _extract_pdf_text(file_path):
    """从PDF文件提取文字（三级回退）"""
    text = ''
    # 方法1: PyMuPDF
    try:
        import fitz
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text
    except Exception as e:
        print(f"PyMuPDF提取失败: {e}")
    
    # 方法2: PyPDF2
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ''
        if text.strip():
            return text
    except Exception as e:
        print(f"PyPDF2提取失败: {e}")
    
    # 方法3: OCR
    try:
        import fitz
        import pytesseract
        from PIL import Image
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        doc = fitz.open(file_path)
        for page_num in range(min(5, len(doc))):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text += pytesseract.image_to_string(img, lang='chi_sim+eng')
        doc.close()
    except Exception as e:
        print(f"OCR提取失败: {e}")
    
    return text

@app.route('/api/gaokao/search')
@login_required
def api_gaokao_search():
    """搜索高考真题API"""
    subject = request.args.get('subject', '')
    knowledge_point = request.args.get('knowledge_point', '')
    question_type = request.args.get('question_type', '')
    category = request.args.get('category', '')
    limit = request.args.get('limit', 5, type=int)
    
    questions = db.search_gaokao_questions(
        subject=subject if subject else None,
        knowledge_point=knowledge_point if knowledge_point else None,
        question_type=question_type if question_type else None,
        category=category if category else None,
        limit=limit
    )
    return jsonify({'success': True, 'questions': questions})

# ==================== 静态文件 ====================
@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# ==================== 启动 ====================
if __name__ == '__main__':
    print("=" * 60)
    print("  高中生试卷管理系统")
    print("  访问地址: http://localhost:5001")
    print("  默认管理员: admin / admin123")
    print("=" * 60)
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5001)), debug=False)
